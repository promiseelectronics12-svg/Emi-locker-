from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "investor"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "EMI_Locker_Three_Tier_Service_Guide.docx"

OUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)


COLORS = {
    "ink": "1E293B",
    "muted": "64748B",
    "line": "CBD5E1",
    "bg": "F8FAFC",
    "normal": "2563EB",
    "premium": "059669",
    "vip": "B45309",
    "normal_light": "DBEAFE",
    "premium_light": "D1FAE5",
    "vip_light": "FEF3C7",
    "dark": "0F172A",
    "white": "FFFFFF",
}


def font(size=32, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def hex_to_rgb(value):
    value = value.strip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def rounded(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def wrap_text(draw, text, font_obj, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textlength(trial, font=font_obj) <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped(draw, text, x, y, font_obj, fill, max_width, line_gap=8):
    for line in wrap_text(draw, text, font_obj, max_width):
        draw.text((x, y), line, font=font_obj, fill=fill)
        y += font_obj.size + line_gap
    return y


def make_tier_cards():
    width, height = 1800, 920
    img = Image.new("RGB", (width, height), hex_to_rgb(COLORS["bg"]))
    draw = ImageDraw.Draw(img)
    title = font(54, True)
    sub = font(28)
    card_title = font(38, True)
    body = font(24)
    small = font(21)

    draw.text((80, 60), "Three Service Tiers", font=title, fill=hex_to_rgb(COLORS["dark"]))
    draw.text(
        (80, 128),
        "Each key tier changes the backend policy, dealer workflow, customer experience, and payment handling.",
        font=sub,
        fill=hex_to_rgb(COLORS["muted"]),
    )

    cards = [
        (
            "Normal",
            "standard",
            COLORS["normal"],
            COLORS["normal_light"],
            "Manual dealer-controlled EMI protection",
            [
                "Dealer receives overdue alerts",
                "Dealer manually locks and unlocks",
                "Basic customer reminders",
                "No payment app required",
            ],
        ),
        (
            "Premium",
            "premium",
            COLORS["premium"],
            COLORS["premium_light"],
            "Backend-assisted automation with dealer payment control",
            [
                "Auto reminders and overdue rules",
                "Dealer confirms payment",
                "Auto lock ladder after grace",
                "Offline rescue remains available",
            ],
        ),
        (
            "VIP",
            "vip",
            COLORS["vip"],
            COLORS["vip_light"],
            "Payment-ledger automation with customer payment app",
            [
                "Separate customer payment app",
                "Advance payment and receipts",
                "Auto relax or unlock after verified payment",
                "Dealer gets payment notifications",
            ],
        ),
    ]
    x_positions = [80, 650, 1220]
    for x, (name, internal, color, light, promise, bullets) in zip(x_positions, cards):
        rounded(draw, (x, 210, x + 500, 820), 28, hex_to_rgb(COLORS["white"]), hex_to_rgb(COLORS["line"]), 3)
        rounded(draw, (x + 28, 238, x + 472, 316), 22, hex_to_rgb(light))
        draw.text((x + 54, 252), name, font=card_title, fill=hex_to_rgb(color))
        draw.text((x + 310, 264), internal, font=small, fill=hex_to_rgb(COLORS["muted"]))
        y = draw_wrapped(draw, promise, x + 38, 360, body, hex_to_rgb(COLORS["dark"]), 420, 10)
        y += 18
        for bullet in bullets:
            draw.ellipse((x + 42, y + 8, x + 56, y + 22), fill=hex_to_rgb(color))
            y = draw_wrapped(draw, bullet, x + 74, y, small, hex_to_rgb(COLORS["ink"]), 380, 7)
            y += 10

    out = ASSET_DIR / "tier_cards.png"
    img.save(out, quality=95)
    return out


def make_flow_diagram():
    width, height = 2400, 980
    img = Image.new("RGB", (width, height), hex_to_rgb(COLORS["white"]))
    draw = ImageDraw.Draw(img)
    title = font(58, True)
    subtitle = font(30)
    label = font(30, True)
    body = font(25)
    draw.text((90, 70), "How a Key Becomes Runtime Policy", font=title, fill=hex_to_rgb(COLORS["dark"]))
    draw.text(
        (90, 140),
        "The selected key tier becomes the active service policy after the device confirms enrollment.",
        font=subtitle,
        fill=hex_to_rgb(COLORS["muted"]),
    )

    steps = [
        ("Activation Key", "standard / premium / vip", COLORS["normal"]),
        ("Pending Enrollment", "dealer creates binding", COLORS["premium"]),
        ("Device Confirmation", "user app confirms code", COLORS["vip"]),
        ("Active EMI Schedule", "service_tier copied here", COLORS["normal"]),
        ("Runtime Policy", "scheduler + app behavior", COLORS["premium"]),
        ("Customer Experience", "user/payment/dealer updates", COLORS["vip"]),
    ]

    # Two-row layout prevents the Word-scaled image from squeezing labels.
    positions = [
        (120, 270),
        (560, 270),
        (1000, 270),
        (1000, 560),
        (560, 560),
        (120, 560),
    ]
    box_w, box_h = 330, 170
    for idx, (head, desc, color) in enumerate(steps):
        x, y = positions[idx]
        rounded(draw, (x, y, x + box_w, y + box_h), 22, hex_to_rgb("F8FAFC"), hex_to_rgb(color), 4)
        draw.text((x + 26, y + 30), head, font=label, fill=hex_to_rgb(color))
        draw_wrapped(draw, desc, x + 26, y + 88, body, hex_to_rgb(COLORS["ink"]), box_w - 52)
        if idx < len(steps) - 1:
            nx, ny = positions[idx + 1]
            start = (x + box_w, y + box_h // 2)
            end = (nx, ny + box_h // 2)
            if y == ny:
                draw.line((start[0] + 12, start[1], end[0] - 12, end[1]), fill=hex_to_rgb(COLORS["line"]), width=6)
                direction = 1 if end[0] > start[0] else -1
                tip_x = end[0] - 12 * direction
                draw.polygon(
                    [
                        (tip_x, end[1]),
                        (tip_x - 22 * direction, end[1] - 14),
                        (tip_x - 22 * direction, end[1] + 14),
                    ],
                    fill=hex_to_rgb(COLORS["line"]),
                )
            else:
                mid_x = x + box_w + 95
                draw.line((start[0] + 12, start[1], mid_x, start[1], mid_x, end[1], end[0] + box_w + 12, end[1]), fill=hex_to_rgb(COLORS["line"]), width=6)
                draw.polygon(
                    [(end[0] + box_w + 12, end[1]), (end[0] + box_w + 34, end[1] - 14), (end[0] + box_w + 34, end[1] + 14)],
                    fill=hex_to_rgb(COLORS["line"]),
                )

    draw_wrapped(
        draw,
        "Developer rule: never mix inventory tier, customer credit tier, and service behavior tier.",
        1500,
        320,
        font(36, True),
        hex_to_rgb(COLORS["dark"]),
        760,
    )
    draw_wrapped(
        draw,
        "The production behavior must use service_tier: Normal = manual, Premium = backend-assisted, VIP = payment-ledger automation.",
        1500,
        450,
        font(29),
        hex_to_rgb(COLORS["muted"]),
        760,
    )
    out = ASSET_DIR / "tier_policy_flow.png"
    img.save(out, quality=95)
    return out


def make_heartbeat_chart():
    width, height = 2400, 1520
    img = Image.new("RGB", (width, height), hex_to_rgb(COLORS["white"]))
    draw = ImageDraw.Draw(img)
    title = font(58, True)
    axis = font(28)
    label = font(32, True)
    value_font = font(27, True)
    small = font(24)
    draw.text((90, 70), "Heartbeat Load Projection", font=title, fill=hex_to_rgb(COLORS["dark"]))
    draw.text(
        (90, 145),
        "Healthy devices heartbeat every 60 minutes. Degraded devices temporarily report every 15 minutes.",
        font=axis,
        fill=hex_to_rgb(COLORS["muted"]),
    )

    groups = ["1k devices", "10k devices", "100k devices"]
    hourly = [0.72, 7.2, 72]
    frequent = [2.88, 28.8, 288]
    max_val = 300

    # Horizontal grouped bars avoid overlap between small 1k labels and large 100k bars.
    chart_x, chart_y = 500, 310
    chart_w = 1500
    row_gap = 290
    bar_h = 48

    draw.text((chart_x, chart_y - 70), "Monthly outbound estimate at about 1 KB per response", font=label, fill=hex_to_rgb(COLORS["dark"]))

    # Axis
    axis_y = chart_y + row_gap * 3 - 15
    draw.line((chart_x, axis_y, chart_x + chart_w, axis_y), fill=hex_to_rgb(COLORS["line"]), width=3)
    for tick in [0, 50, 100, 150, 200, 250, 300]:
        x = chart_x + int((tick / max_val) * chart_w)
        draw.line((x, chart_y - 15, x, axis_y + 8), fill=hex_to_rgb("E2E8F0"), width=1)
        if tick not in (0, 300):
            draw.text((x - 22, axis_y + 22), str(tick), font=small, fill=hex_to_rgb(COLORS["muted"]))
    draw.text((chart_x, axis_y + 22), "Scale: GB/month", font=small, fill=hex_to_rgb(COLORS["muted"]))

    for i, group in enumerate(groups):
        y = chart_y + i * row_gap
        draw.text((90, y + 42), group, font=label, fill=hex_to_rgb(COLORS["dark"]))

        hourly_w = max(7, int((hourly[i] / max_val) * chart_w))
        frequent_w = max(7, int((frequent[i] / max_val) * chart_w))

        # Row backgrounds
        rounded(draw, (chart_x, y + 20, chart_x + chart_w, y + 20 + bar_h), 14, hex_to_rgb("F1F5F9"))
        rounded(draw, (chart_x, y + 102, chart_x + chart_w, y + 102 + bar_h), 14, hex_to_rgb("F1F5F9"))

        rounded(draw, (chart_x, y + 20, chart_x + hourly_w, y + 20 + bar_h), 14, hex_to_rgb(COLORS["premium"]))
        rounded(draw, (chart_x, y + 102, chart_x + frequent_w, y + 102 + bar_h), 14, hex_to_rgb(COLORS["vip"]))

        draw.text((chart_x - 225, y + 27), "Hourly", font=axis, fill=hex_to_rgb(COLORS["premium"]))
        draw.text((chart_x - 225, y + 109), "15-minute", font=axis, fill=hex_to_rgb(COLORS["vip"]))

        hourly_label = f"{hourly[i]} GB / month"
        frequent_label = f"{frequent[i]} GB / month"

        if hourly_w > 620:
            draw.text((chart_x + hourly_w - 385, y + 26), hourly_label, font=value_font, fill=hex_to_rgb(COLORS["white"]))
        else:
            h_label_x = max(chart_x + hourly_w + 24, chart_x + 135)
            draw.text((h_label_x, y + 26), hourly_label, font=value_font, fill=hex_to_rgb(COLORS["premium"]))

        if frequent_w > 620:
            draw.text((chart_x + frequent_w - 400, y + 108), frequent_label, font=value_font, fill=hex_to_rgb(COLORS["white"]))
        else:
            f_label_x = max(chart_x + frequent_w + 24, chart_x + 135)
            draw.text((f_label_x, y + 108), frequent_label, font=value_font, fill=hex_to_rgb(COLORS["vip"]))

        draw.text(
            (chart_x, y + 168),
            f"Requests: {['720k', '7.2M', '72M'][i]} monthly hourly vs {['2.88M', '28.8M', '288M'][i]} monthly at 15-minute degraded mode.",
            font=small,
            fill=hex_to_rgb(COLORS["muted"]),
        )

    rounded(draw, (90, 1320, 2180, 1415), 24, hex_to_rgb("FEF3C7"), hex_to_rgb("F59E0B"), 3)
    draw.text((125, 1348), "Conclusion: 1,000 devices are manageable with hourly heartbeat. Production scale needs paid hosting, indexes, batching, and observability.", font=axis, fill=hex_to_rgb(COLORS["dark"]))
    out = ASSET_DIR / "heartbeat_load_projection.png"
    img.save(out, quality=95)
    return out


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="1E293B", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor.from_string(COLORS["dark"])
    return p


def add_note_box(doc, title, body, fill="F8FAFC"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(COLORS["dark"])
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(4)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True, COLORS["white"], 8)
        shade_cell(hdr[i], COLORS["dark"])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), False, COLORS["ink"], 8)
            shade_cell(cells[i], "FFFFFF")
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def build_docx():
    tier_cards = make_tier_cards()
    flow = make_flow_diagram()
    heartbeat = make_heartbeat_chart()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Heading 1", 20), ("Heading 2", 15), ("Heading 3", 12)]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True
        styles[style_name].font.color.rgb = RGBColor.from_string(COLORS["dark"])

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("EMI Locker")
    r.bold = True
    r.font.size = Pt(32)
    r.font.color.rgb = RGBColor.from_string(COLORS["dark"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Three-Tier Service Guide")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(COLORS["normal"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Investor-facing service separation for Normal, Premium, and VIP EMI protection")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    doc.add_picture(str(tier_cards), width=Inches(7.3))
    doc.add_paragraph()
    add_note_box(
        doc,
        "Product Position",
        "EMI Locker protects both parties: dealer confidence without customer humiliation. The platform starts with reminders, then disturbance, then restriction, and uses full lock only when business rules justify it.",
        "EFF6FF",
    )
    doc.add_page_break()

    add_heading(doc, "Executive Summary", 1)
    doc.add_paragraph(
        "The three tiers are not just pricing labels. Each tier defines a different operating model for payment handling, backend automation, customer experience, and dealer workload."
    )
    add_table(
        doc,
        ["Tier", "Best For", "Payment Handling", "Lock Behavior", "Dealer Workload"],
        [
            ["Normal", "Small shops", "Manual by dealer", "Dealer manual lock/unlock", "High"],
            ["Premium", "Serious dealers", "Dealer confirms payment", "Backend reminders + auto lock ladder", "Medium"],
            ["VIP", "Large partners", "Verified payment ledger + payment app", "Auto relax/unlock after verified payment", "Low"],
        ],
        [1.0, 1.35, 1.75, 2.1, 1.1],
    )

    add_heading(doc, "Runtime Policy Flow", 1)
    doc.add_picture(str(flow), width=Inches(7.25))
    add_note_box(
        doc,
        "Implementation Rule",
        "Key tier must become service_tier during confirmed enrollment. The scheduler, dealer app, user app, and VIP payment app must all read service_tier behavior.",
        "F8FAFC",
    )

    doc.add_page_break()
    add_heading(doc, "Normal Tier", 1)
    doc.add_paragraph("Normal is the low-cost, dealer-operated tier. The platform gives the dealer tools, but the dealer makes the business decisions.")
    add_bullets(
        doc,
        [
            "Backend stores EMI schedule and notifies the dealer about due or overdue state.",
            "Backend does not auto-lock by default.",
            "Dealer manually locks and unlocks from the dealer app.",
            "User app acts as the core managed device agent and reports heartbeat, lock state, permissions, and location when requested.",
            "No separate payment app is required.",
        ],
    )

    add_heading(doc, "Premium Tier", 1)
    doc.add_paragraph("Premium reduces dealer workload. The backend watches EMI dates and applies rules automatically, but payment confirmation remains dealer-controlled.")
    add_bullets(
        doc,
        [
            "Backend sends reminders before due dates and alerts at overdue state.",
            "Backend applies reminder mode, partial lock, and full lock based on grace/overdue rules.",
            "Dealer records or confirms payments and controls unlock or grace decisions.",
            "Dealer app shows automation timeline: reminder sent, warning shown, partial/full lock applied, device confirmed.",
            "User app explains EMI status, lock reason, due date, grace period, and dealer contact.",
        ],
    )

    add_heading(doc, "VIP Tier", 1)
    doc.add_paragraph("VIP is the PalmPay-style experience: payment becomes part of the customer journey, while the dealer receives payment notifications and exception alerts.")
    add_bullets(
        doc,
        [
            "VIP devices receive a separate customer payment app or portal in addition to the core user app.",
            "Payment ledger tracks due, pending, verified, rejected, refunded, and dealer settlement states.",
            "Customer can see balance, installments left, pay-now, advance payment, receipt history, and support contact.",
            "Verified payment automatically relaxes or unlocks the device when business rules allow.",
            "Full lock is not the first response. VIP uses persistent disturbance mode first, keeping payment/support paths usable.",
        ],
    )

    add_heading(doc, "VIP Lock Ladder", 2)
    add_table(
        doc,
        ["Stage", "Trigger", "Customer Impact"],
        [
            ["Friendly reminder", "Before due date", "Notification and payment-app prompt"],
            ["Payment due", "Due date or grace start", "Persistent payment card in payment/user app"],
            ["Disturbance mode", "After grace if unpaid", "Front message that obstructs view but keeps payment/support usable"],
            ["Partial restriction", "Continued overdue or tamper", "Restrict non-essential/distraction apps when policy allows"],
            ["Full lock", "Serious default or fraud", "Kiosk/lock mode with emergency, payment, and dealer contact only"],
        ],
        [1.4, 1.8, 3.4],
    )

    doc.add_page_break()
    add_heading(doc, "App Separation", 1)
    add_note_box(
        doc,
        "Decision",
        "Use one core user app for all tiers, plus one optional VIP payment app. Do not create separate Normal, Premium, and VIP user apps.",
        "ECFDF5",
    )
    add_table(
        doc,
        ["App", "Installed On", "Responsibility"],
        [
            ["Core user app", "Every EMI-managed phone", "Device Owner/admin behavior, FCM commands, lock overlay/kiosk, heartbeat, schedule sync, offline unlock, location report"],
            ["Dealer app", "Dealer/reseller device", "Enrollment, key purchase, device status, lock/unlock, payment confirmation, notifications"],
            ["VIP payment app", "VIP devices only", "Customer payment UX, payment initiation, payment status, receipt history, advance payment, support"],
        ],
        [1.3, 1.4, 4.5],
    )
    doc.add_paragraph(
        "For Device Owner / AMAPI-managed phones, the backend can force-install the VIP payment app through managed Google Play policy using the app package name."
    )

    add_heading(doc, "Backend Readiness Work", 1)
    add_table(
        doc,
        ["Workstream", "Required Change"],
        [
            ["Source of truth", "Add/copy service_tier into active device or schedule during confirmed enrollment"],
            ["Tier-aware scheduler", "standard = notify only; premium = auto lock ladder; vip = payment-ledger automation"],
            ["Payment ledger", "Normalize payment states and add idempotent payment events/webhooks"],
            ["Command lifecycle", "Track created, sent, received, applied, failed, expired"],
            ["Dealer UI", "Show tier detail sheet, tier badge, tier-specific actions, and command timeline"],
            ["User/VIP app", "Read tier from backend; VIP opens/links payment app and applies disturbance mode first"],
        ],
        [1.6, 5.2],
    )

    doc.add_page_break()
    add_heading(doc, "Heartbeat And Scaling", 1)
    doc.add_picture(str(heartbeat), width=Inches(7.25))
    doc.add_paragraph(
        "Heartbeat is necessary, but it must be controlled. Healthy devices should report about once per hour. Permission-degraded or recently changed devices can report more often temporarily."
    )
    add_table(
        doc,
        ["Devices", "Hourly Heartbeat", "15-Minute Degraded Heartbeat"],
        [
            ["1,000", "720,000/month, about 0.7 GB outbound at 1 KB", "2.88M/month, about 2.9 GB outbound at 1 KB"],
            ["10,000", "7.2M/month, about 7.2 GB outbound at 1 KB", "28.8M/month, about 28.8 GB outbound at 1 KB"],
            ["100,000", "72M/month, about 72 GB outbound at 1 KB", "288M/month, about 288 GB outbound at 1 KB"],
        ],
        [1.0, 2.8, 3.3],
    )
    add_note_box(
        doc,
        "Infrastructure Note",
        "Render free tier is acceptable for development testing only. Production needs paid hosting, database indexes, batching, and observability.",
        "FEF3C7",
    )

    add_heading(doc, "References", 1)
    refs = [
        "Android Management API: Distribute apps through managed Google Play policy - https://developers.google.com/android/management/apps",
        "Render free tier limitations - https://render.com/docs/free",
        "Render outbound bandwidth pricing - https://render.com/docs/outbound-bandwidth",
    ]
    add_bullets(doc, refs)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "EMI Locker Three-Tier Service Guide | Confidential investor draft"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_docx()
    print(path)



