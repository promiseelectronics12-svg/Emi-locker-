#!/usr/bin/env python3
"""
EMI Locker Platform — Blueprint & Developer Debug Guide Generator
Produces: EMI_Locker_Blueprint.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette ────────────────────────────────────────────────────────────
C_BLACK      = RGBColor(0x11, 0x18, 0x27)
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_BLUE       = RGBColor(0x25, 0x63, 0xEB)
C_BLUE_LT    = RGBColor(0xEF, 0xF6, 0xFF)
C_GREEN      = RGBColor(0x05, 0x96, 0x69)
C_GREEN_LT   = RGBColor(0xEC, 0xFD, 0xF5)
C_RED        = RGBColor(0xDC, 0x26, 0x26)
C_RED_LT     = RGBColor(0xFE, 0xF2, 0xF2)
C_AMBER      = RGBColor(0xB4, 0x53, 0x09)
C_AMBER_LT   = RGBColor(0xFF, 0xFB, 0xEB)
C_VIOLET     = RGBColor(0x7C, 0x3A, 0xED)
C_VIOLET_LT  = RGBColor(0xF5, 0xF3, 0xFF)
C_CYAN       = RGBColor(0x08, 0x91, 0xB2)
C_CYAN_LT    = RGBColor(0xEC, 0xFE, 0xFF)
C_GREY       = RGBColor(0x64, 0x74, 0x8B)
C_GREY_LT    = RGBColor(0xF8, 0xFA, 0xFC)
C_PANEL      = RGBColor(0xF1, 0xF5, 0xF9)

def hex_color(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(rgb))
    tcPr.append(shd)

def set_cell_borders(cell, color="D1D5DB"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color or C_BLACK
        if level == 1:
            run.font.size = Pt(20)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(15)
            run.font.bold = True
        else:
            run.font.size = Pt(12)
            run.font.bold = True
    return p

def add_para(doc, text, bold=False, italic=False, color=None, size=10, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or C_BLACK
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bottom)
    pPr.append(pBdr)

def make_table(doc, headers, rows, col_widths=None, header_bg=None, stripe=True):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg or C_BLACK)
        set_cell_borders(cell, "374151")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_WHITE

    # Data rows
    for r_i, row in enumerate(rows):
        tr = table.rows[r_i + 1]
        bg = C_GREY_LT if (stripe and r_i % 2 == 1) else C_WHITE
        for c_i, val in enumerate(row):
            cell = tr.cells[c_i]
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if isinstance(val, tuple):
                text, color, bold = val
                run = p.add_run(str(text))
                run.font.color.rgb = color
                run.font.bold = bold
                run.font.size = Pt(9)
            else:
                run = p.add_run(str(val))
                run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── COVER PAGE ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
run = p.add_run("EMI LOCKER PLATFORM")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = C_BLACK
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run("Developer Blueprint & Debug Reference")
run.font.size = Pt(16)
run.font.color.rgb = C_GREY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
run = p.add_run("Bangladesh Android MDM · EMI Phone Financing Security Platform")
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = C_GREY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}  ·  Version 1.0")
run.font.size = Pt(10)
run.font.color.rgb = C_GREY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── SECTION 1: SYSTEM OVERVIEW ───────────────────────────────────────────────
add_heading(doc, "1. System Overview", 1, C_BLUE)
add_para(doc,
    "EMI Locker is an Android MDM (Mobile Device Management) platform for phone dealers "
    "in Bangladesh who sell phones on installment (EMI). When a customer misses a payment, "
    "the server progressively locks the device through Android Management API (AMAPI). "
    "When the final payment is confirmed by an admin, the device is fully released.", size=10)
doc.add_paragraph()

add_heading(doc, "1.1  Stakeholder Hierarchy", 2)
make_table(doc,
    ["Role", "Access Level", "Primary Actions"],
    [
        ["Admin",    "Full system access + 2FA required for all critical actions",
                     "Approve resellers · Execute decoupling · Review audit log · Override locks"],
        ["Reseller", "Manages dealers, requests activation key quota from admin",
                     "Request keys · Assign keys to dealers · View dealer performance"],
        ["Dealer",   "Manages devices and customers, initiates lock/unlock requests",
                     "Enroll devices · Request lock · View EMI · Pull GPS · Flag fraud"],
        ["Customer", "Read-only via Kotlin app on their device",
                     "View EMI schedule · Receive lock notifications · Use PAUT offline token"],
    ],
    col_widths=[1.1, 2.4, 3.0], header_bg=C_BLUE
)

doc.add_paragraph()
add_heading(doc, "1.2  8-Layer Security Model", 2)
make_table(doc,
    ["Layer", "Name", "What It Does", "Bypass Consequence"],
    [
        ["L1", "Device Owner (DPC)", "AMAPI enrolls phone as managed — persists through factory reset", (  "Full control lost", C_RED, True)],
        ["L2", "FRP", "Factory Reset Protection via managed Google account — blocks re-use after wipe", ("Device resold freely", C_RED, True)],
        ["L3", "Managed Google Account", "Google account tied to device; removed only on admin-executed decouple", ("FRP clears on its own", C_RED, True)],
        ["L4", "Certificate Pinning", "Kotlin app pins TLS cert — dev mode disables when localhost detected", ("MITM possible in prod", C_AMBER, True)],
        ["L5", "HSM Command Signing", "All lock/unlock/decouple commands signed via KMS before dispatch", ("Unsigned commands accepted", C_RED, True)],
        ["L6", "Hardware Binding", "IMEI+Serial+SoC fingerprint bound at enrollment — rejects cloned devices", ("Cloned phone unlocks", C_RED, True)],
        ["L7", "APK Integrity", "Play Integrity API attestation — auto-lock on tampered APK", ("Modded app runs freely", C_AMBER, True)],
        ["L8", "Samsung Knox", "Additional policy enforcement on Samsung devices", ("Knox bypass on Samsung", C_AMBER, False)],
    ],
    col_widths=[0.4, 1.6, 2.8, 1.8], header_bg=C_BLACK
)
doc.add_page_break()

# ── SECTION 2: ARCHITECTURE DIAGRAM ──────────────────────────────────────────
add_heading(doc, "2. Architecture Diagram", 1, C_BLUE)
add_para(doc, "Request flow from client to database to external services:", size=10)
doc.add_paragraph()

arch_lines = [
    "  ┌──────────────────┐  ┌──────────────────┐  ┌──────────────────┐",
    "  │  React Admin     │  │  Flutter Dealer  │  │  Kotlin User App │",
    "  │  Panel (:5173)   │  │  App (:3000 API) │  │  (:3000 API)     │",
    "  └────────┬─────────┘  └────────┬─────────┘  └────────┬─────────┘",
    "           │                     │                      │",
    "           └─────────────────────┼──────────────────────┘",
    "                                 │  HTTPS / JWT (RS256)",
    "                                 ▼",
    "  ┌──────────────────────────────────────────────────────────────────┐",
    "  │              Node.js + Express API  (Port 3000)                  │",
    "  │  /api/v1/auth  /lock  /emi  /devices  /keys  /decoupling  ...   │",
    "  │                                                                  │",
    "  │  Global Middleware:  helmet · cors · express.json · morgan       │",
    "  │  Auth Middleware:    requireAuth (RS256 JWT) · requireRole()     │",
    "  │  Error Handler:      errorHandler.js  (registered LAST)         │",
    "  └───────┬──────────────────────────────────────────┬──────────────┘",
    "          │                                          │",
    "   ┌──────▼──────┐                          ┌───────▼──────┐",
    "   │  PostgreSQL │                          │    Redis     │",
    "   │  (AES-256   │                          │  Sessions    │",
    "   │  at rest)   │                          │  Rate limits │",
    "   └─────────────┘                          │  Token BL    │",
    "                                            └──────────────┘",
    "          │ External Services",
    "   ┌──────▼──────┐  ┌──────────────┐  ┌──────────────┐  ┌────────────┐",
    "   │  Firebase   │  │ Google AMAPI │  │   Twilio SMS │  │  KMS / HSM │",
    "   │  FCM Push   │  │ Device Mgmt  │  │  (fallback)  │  │  Signing   │",
    "   └─────────────┘  └──────────────┘  └──────────────┘  └────────────┘",
]
for line in arch_lines:
    add_code(doc, line)

doc.add_paragraph()
add_heading(doc, "2.1  Lock Delivery — 3-Channel Flow", 2)
lock_lines = [
    "  Trigger: Payment overdue / dealer request / admin override",
    "       │",
    "       ▼",
    "  ┌─────────────────────┐",
    "  │   LockService       │  ← Entry point. NEVER bypass this.",
    "  │  .requestLock()     │    All lock actions go through here.",
    "  └──────────┬──────────┘",
    "             │",
    "       ┌─────▼──────────────────────────────────────────────────┐",
    "       │              LockDeliveryService                       │",
    "       │                                                        │",
    "       │  Channel 1: FCM  ──────► FirebaseService.send()       │",
    "       │     │ (fails?)                                        │",
    "       │  Channel 2: AMAPI ─────► AmapiService.applyPolicy()  │",
    "       │     │ (fails?)                                        │",
    "       │  Channel 3: PAUT/PADT ─► PautService.issue()         │",
    "       │             (offline token, device checks on reconnect)│",
    "       └────────────────────────────────────────────────────────┘",
    "             │",
    "  LockVerificationService: waits 60s, confirms device ACK.",
    "  If no ACK: escalates to Channel 2 or 3.",
]
for line in lock_lines:
    add_code(doc, line)

doc.add_paragraph()
add_heading(doc, "2.2  Decoupling State Machine", 2)
sm_lines = [
    "  EMI_ACTIVE",
    "       │  (final payment confirmed)",
    "       ▼",
    "  FINAL_PAYMENT_RECEIVED",
    "       │  (auto: notify dealer, start 5-day Bull queue timer)",
    "       ▼",
    "  DEALER_NOTIFIED",
    "       │                    │",
    "       │ (no fraud flag)    │ (dealer flags fraud — does NOT block decouple)",
    "       ▼                    ▼",
    "  PENDING_ADMIN_DECOUPLE   FRAUD_FLAGGED",
    "       │                    │  (admin reviews but executes regardless)",
    "       └──────────┬─────────┘",
    "                  │  (admin executes — requires 2FA)",
    "                  ▼",
    "         Generate RTOC → Send signed FCM command → AmapiService.deleteManagedAccount()",
    "                  │",
    "          ┌───────▼───────┐    ┌─────────────────────────────────┐",
    "          │DEVICE_DECOUPLED│    │ FCM failed? Issue PADT (7 days) │",
    "          │  (immutable   │    │ device checks on reconnect       │",
    "          │   audit log)  │    └─────────────────────────────────┘",
    "          └───────────────┘",
]
for line in sm_lines:
    add_code(doc, line)

doc.add_page_break()

# ── SECTION 3: MODULE MAP ─────────────────────────────────────────────────────
add_heading(doc, "3. Module Map", 1, C_BLUE)
add_para(doc, "Every module in the backend, its location, files, and current build status.", size=10)
doc.add_paragraph()

STATUS_OK   = (  "✓ APPROVED", C_GREEN, True)
STATUS_WARN = ("⚠ RE-AUDIT",  C_AMBER, True)
STATUS_FIX  = ("↻ FIXING",    C_CYAN,  True)
STATUS_FAIL = ("✗ FAILED",    C_RED,   True)

make_table(doc,
    ["Module ID", "Path", "Key Files", "Phase", "Status", "Last Score"],
    [
        ["project-structure",   "backend/src · user-app · dealer-app · admin-panel · database",
         "package.json · pubspec.yaml · build.gradle", "1", STATUS_OK, "—"],
        ["database-schema",     "database/migrations/",
         "001_users.sql → 010_paut_tokens.sql · seeds/roles.sql", "1", STATUS_FIX, "82"],
        ["backend-auth",        "backend/src/modules/auth/",
         "jwt.js · totp.js · rateLimit.js · index.js", "1", STATUS_FIX, "62"],
        ["backend-devices",     "backend/src/modules/devices/",
         "deviceService.js · amapiService.js · commandSigningService.js · hardwareBindingService.js · kmsSigningService.js", "1", STATUS_OK, "95"],
        ["backend-lock-engine", "backend/src/modules/lock/",
         "lockService.js · lockDeliveryService.js · lockCommandService.js · lockSchedulerService.js · pautService.js · padtService.js", "1", STATUS_FIX, "92"],
        ["backend-emi",         "backend/src/modules/emi/",
         "emiService.js · emiController.js · emiModel.js · emiValidation.js · decouplingService.js", "1", STATUS_FIX, "75"],
        ["backend-notifications","backend/src/modules/notifications/",
         "fcm.service.js · notification.service.js · sms.service.js · dealer-message-rate-limiter.js", "1", STATUS_FIX, "82"],
        ["backend-server-entry","backend/src/index.js · backend/src/config/",
         "index.js · database.js · redis.js · envValidator.js · errorHandler.js", "1", STATUS_FIX, "72"],
        ["backend-keys",        "backend/src/modules/keys/",
         "keyService.js · hmac.service.js · keyScheduler.js", "2", STATUS_FIX, "TBC"],
        ["backend-decoupling",  "backend/src/modules/decoupling/",
         "decouplingService.js · decouplingScheduler.js · decouplingModel.js", "2", STATUS_FIX, "75"],
        ["backend-location",    "backend/src/modules/location/",
         "locationService.js · locationScheduler.js", "2", STATUS_FIX, "72"],
        ["backend-admin-api",   "backend/src/modules/admin/",
         "adminService.js · adminDeviceService.js · adminController.js · adminRoutes.js", "2", STATUS_FIX, "78"],
        ["backend-fraud",       "backend/src/modules/fraud/",
         "fraudService.js · fraudScheduler.js · fraudController.js", "2", STATUS_FIX, "75"],
        ["user-app-foundation", "user-app/",
         "build.gradle · DpcReceiver.kt · MainActivity.kt", "3", STATUS_OK, "—"],
    ],
    col_widths=[1.4, 1.6, 2.2, 0.5, 0.85, 0.75], header_bg=C_BLUE
)
doc.add_paragraph()
add_para(doc, "✓ APPROVED = meets 95+ standard  ·  ↻ FIXING = Gemini fixer running  ·  ⚠ RE-AUDIT = built under old threshold  ·  ✗ FAILED = needs rebuild", italic=True, color=C_GREY, size=9)

doc.add_page_break()

# ── SECTION 4: ERROR CODE TAXONOMY ───────────────────────────────────────────
add_heading(doc, "4. Error Code Taxonomy", 1, C_BLUE)
add_para(doc,
    "Every error the server returns uses a structured error code. Format: [MODULE]-[CATEGORY][NUMBER]. "
    "These codes appear in API responses, server logs, and the admin panel debug log. "
    "Use this table to locate the exact file and function when a code appears.", size=10)
doc.add_paragraph()

add_heading(doc, "4.1  Error Code Format", 2)
add_code(doc, "  EMI-AUTH-101")
add_code(doc, "      │    │   └── Error number within category")
add_code(doc, "      │    └────── Module prefix (see table below)")
add_code(doc, "      └─────────── Platform prefix (always EMI)")
doc.add_paragraph()

add_heading(doc, "4.2  Category Ranges", 2)
make_table(doc,
    ["Range", "Category", "Examples"],
    [
        ["001–099", "Startup / Runtime",   "Service not initialised, DB connection failed, Redis unavailable"],
        ["100–199", "Auth / Authorization","Invalid token, expired token, insufficient role, 2FA required"],
        ["200–299", "Database",            "Unique constraint, record not found, transaction rollback"],
        ["300–399", "External Services",   "FCM delivery failed, AMAPI error, Twilio SMS failed, KMS error"],
        ["400–499", "Validation",          "Missing required field, invalid ENUM value, format mismatch"],
        ["500–599", "State Machine",       "Invalid state transition, wrong prerequisite state"],
        ["600–699", "Security Violations", "HMAC mismatch, hardware binding fail, APK integrity fail"],
        ["700–799", "Rate Limit / Quota",  "Login attempts exceeded, dealer message limit, key velocity alert"],
    ],
    col_widths=[0.9, 1.5, 4.1], header_bg=C_BLACK
)
doc.add_paragraph()

add_heading(doc, "4.3  Complete Error Code Registry", 2)
add_para(doc, "Add these error codes to every thrown error in the codebase: throw createError('EMI-AUTH-101', 'Token expired')", size=9, italic=True, color=C_GREY)
doc.add_paragraph()

error_codes = [
    # AUTH
    ("EMI-AUTH-001", "AUTH", "backend/src/modules/auth/jwt.js",           "verifyToken()",         "JWT verification failed — malformed token"),
    ("EMI-AUTH-002", "AUTH", "backend/src/modules/auth/jwt.js",           "verifyToken()",         "JWT signature invalid — possible tampering"),
    ("EMI-AUTH-101", "AUTH", "backend/src/modules/auth/jwt.js",           "verifyToken()",         "Access token expired — client must refresh"),
    ("EMI-AUTH-102", "AUTH", "backend/src/modules/auth/jwt.js",           "verifyToken()",         "Refresh token expired — user must re-login"),
    ("EMI-AUTH-103", "AUTH", "backend/src/middleware/auth.js",            "requireAuth()",         "Token blacklisted — already logged out"),
    ("EMI-AUTH-104", "AUTH", "backend/src/modules/auth/totp.js",          "verify()",              "TOTP code invalid or expired"),
    ("EMI-AUTH-105", "AUTH", "backend/src/middleware/rbac.js",            "requireRole()",         "Insufficient role — action requires higher privilege"),
    ("EMI-AUTH-106", "AUTH", "backend/src/modules/auth/rateLimit.js",     "loginLimiter",          "Too many login attempts — IP blocked for 15 minutes"),
    ("EMI-AUTH-700", "AUTH", "backend/src/modules/auth/rateLimit.js",     "loginLimiter",          "Rate limit exceeded — 429 returned to client"),
    # DEVICES
    ("EMI-DEV-001",  "DEV",  "backend/src/modules/devices/deviceService.js",       "enrollDevice()",   "Device enrollment failed — AMAPI registration error"),
    ("EMI-DEV-201",  "DEV",  "backend/src/modules/devices/device.repository.js",   "findByImei()",     "Device IMEI not found in database"),
    ("EMI-DEV-202",  "DEV",  "backend/src/modules/devices/device.repository.js",   "create()",         "IMEI already registered — duplicate enrollment"),
    ("EMI-DEV-301",  "DEV",  "backend/src/modules/devices/amapiService.js",        "enrollDevice()",   "AMAPI API call failed — check AMAPI_PROJECT env var"),
    ("EMI-DEV-601",  "DEV",  "backend/src/modules/devices/hardwareBindingService.js","verify()",       "Hardware fingerprint mismatch — possible cloned device"),
    ("EMI-DEV-602",  "DEV",  "backend/src/modules/devices/kmsSigningService.js",   "sign()",           "KMS signing failed — command not dispatched"),
    # LOCK
    ("EMI-LOCK-001", "LOCK", "backend/src/modules/lock/lockService.js",            "requestLock()",    "Lock service initialisation error"),
    ("EMI-LOCK-301", "LOCK", "backend/src/modules/lock/lockDeliveryService.js",    "deliverViaFcm()",  "FCM Channel 1 delivery failed — trying Channel 2"),
    ("EMI-LOCK-302", "LOCK", "backend/src/modules/lock/lockDeliveryService.js",    "deliverViaAmapi()", "AMAPI Channel 2 delivery failed — issuing PAUT"),
    ("EMI-LOCK-303", "LOCK", "backend/src/modules/lock/lockDeliveryService.js",    "allChannelsFailed()", "All 3 channels failed — PAUT issued, device offline"),
    ("EMI-LOCK-500", "LOCK", "backend/src/modules/lock/lockService.js",            "requestLock()",    "Invalid lock level — must be SOFT|PARTIAL|FULL_LOCK"),
    ("EMI-LOCK-601", "LOCK", "backend/src/modules/lock/lockCommandService.js",     "buildCommand()",   "Command signing failed — HMAC_SECRET may be missing"),
    ("EMI-LOCK-602", "LOCK", "backend/src/modules/lock/pautService.js",            "issue()",          "PAUT token issue failed — check HMAC_SECRET env var"),
    # EMI
    ("EMI-EMI-400",  "EMI",  "backend/src/modules/emi/emiValidation.js",           "validateSchedule()","Invalid EMI schedule — amount or duration out of range"),
    ("EMI-EMI-401",  "EMI",  "backend/src/modules/emi/emiValidation.js",           "validatePayment()", "Duplicate payment — same installment already confirmed"),
    ("EMI-EMI-402",  "EMI",  "backend/src/modules/emi/emiService.js",              "requestGracePeriod()", "Grace period limit reached — max 2 extensions per cycle"),
    ("EMI-EMI-403",  "EMI",  "backend/src/modules/emi/emiService.js",              "requestGracePeriod()", "Grace days exceed maximum — max 14 days per extension"),
    ("EMI-EMI-500",  "EMI",  "backend/src/modules/emi/emiService.js",              "triggerDecoupling()", "Final payment detected but decoupling trigger failed"),
    # NOTIFICATIONS
    ("EMI-NOTIF-301","NOTIF","backend/src/modules/notifications/fcm.service.js",   "sendToDevice()",   "FCM delivery failed — token may be stale"),
    ("EMI-NOTIF-302","NOTIF","backend/src/modules/notifications/sms.service.js",   "sendSms()",        "Twilio SMS delivery failed — check credentials"),
    ("EMI-NOTIF-700","NOTIF","backend/src/modules/notifications/dealer-message-rate-limiter.js","check()", "Dealer message limit reached — max 10 per device per day"),
    # KEYS
    ("EMI-KEY-400",  "KEY",  "backend/src/modules/keys/keyService.js",             "requestKeys()",    "Quantity exceeds 20% of monthly quota per single request"),
    ("EMI-KEY-401",  "KEY",  "backend/src/modules/keys/keyService.js",             "consumeKey()",     "Key already consumed — cannot reuse activation key"),
    ("EMI-KEY-402",  "KEY",  "backend/src/modules/keys/keyService.js",             "consumeKey()",     "Key expired — 72-hour window has passed"),
    ("EMI-KEY-403",  "KEY",  "backend/src/modules/keys/keyService.js",             "consumeKey()",     "Key does not belong to this dealer"),
    ("EMI-KEY-601",  "KEY",  "backend/src/modules/keys/hmac.service.js",           "verify()",         "HMAC signature mismatch — key may be forged"),
    ("EMI-KEY-700",  "KEY",  "backend/src/modules/keys/keyScheduler.js",           "velocityCheck()",  "Key velocity alert — 10+ consumptions in 24h"),
    # DECOUPLING
    ("EMI-DCP-500",  "DCP",  "backend/src/modules/decoupling/decouplingService.js","fraudFlag()",      "Cannot fraud-flag — device not in DEALER_NOTIFIED state"),
    ("EMI-DCP-501",  "DCP",  "backend/src/modules/decoupling/decouplingService.js","executeDecouple()","Cannot execute — device not in PENDING_ADMIN_DECOUPLE state"),
    ("EMI-DCP-302",  "DCP",  "backend/src/modules/decoupling/decouplingService.js","executeDecouple()","AMAPI delete managed account failed — FRP may not be cleared"),
    ("EMI-DCP-303",  "DCP",  "backend/src/modules/decoupling/decouplingService.js","executeDecouple()","FCM decouple command failed — PADT issued as fallback"),
    # LOCATION
    ("EMI-LOC-301",  "LOC",  "backend/src/modules/location/locationService.js",    "pullGPS()",        "FCM GPS pull command failed — device may be offline"),
    ("EMI-LOC-400",  "LOC",  "backend/src/modules/location/locationService.js",    "reportGPS()",      "Invalid GPS coordinates — latitude or longitude out of range"),
    # ADMIN
    ("EMI-ADM-100",  "ADM",  "backend/src/modules/admin/adminMiddleware.js",       "require2FA()",     "Admin action requires verified 2FA session"),
    ("EMI-ADM-200",  "ADM",  "backend/src/modules/admin/adminService.js",          "getDashboard()",   "Dashboard query failed — DB connection error"),
    # FRAUD
    ("EMI-FRD-301",  "FRD",  "backend/src/modules/fraud/fraudService.js",          "handleIntegrityReport()", "Play Integrity webhook signature invalid — rejected"),
    ("EMI-FRD-601",  "FRD",  "backend/src/modules/fraud/fraudService.js",          "handleIntegrityReport()", "APK integrity failed — device auto-locked"),
    # SERVER
    ("EMI-SRV-001",  "SRV",  "backend/src/config/database.js",                    "connectDB()",      "PostgreSQL connection failed at startup — check POSTGRES_URL"),
    ("EMI-SRV-002",  "SRV",  "backend/src/config/redis.js",                       "connectRedis()",   "Redis connection failed at startup — check REDIS_URL"),
    ("EMI-SRV-003",  "SRV",  "backend/src/config/envValidator.js",                "validateEnv()",    "Required environment variable missing — server will not start"),
    # DATABASE
    ("EMI-DB-200",   "DB",   "database/migrations/",                              "all queries",      "Unique constraint violation"),
    ("EMI-DB-201",   "DB",   "database/migrations/",                              "all queries",      "Foreign key constraint violation"),
    ("EMI-DB-202",   "DB",   "database/migrations/008_audit_log.sql",             "audit_log trigger","Attempt to UPDATE/DELETE audit_log — operation blocked by trigger"),
]

make_table(doc,
    ["Error Code", "Module", "File Path", "Function", "Description"],
    [
        (   (code, C_BLUE, True),
            mod,
            path,
            func,
            desc
        )
        for code, mod, path, func, desc in error_codes
    ],
    col_widths=[1.1, 0.55, 2.0, 1.5, 2.35], header_bg=C_BLACK
)

doc.add_page_break()

# ── SECTION 5: KNOWN PROBLEMATIC AREAS ───────────────────────────────────────
add_heading(doc, "5. Known Problematic Areas", 1, C_RED)
add_para(doc, "These are areas where bugs were found during the build process. "
    "A future developer should inspect these files first when debugging.", size=10)
doc.add_paragraph()

problems = [
    {
        "title": "CRITICAL — db.connect() does not exist",
        "severity": "HIGH",
        "modules": "backend-admin-api, any module that imports database.js",
        "files": "backend/src/modules/admin/adminService.js\nbackend/src/modules/admin/adminDeviceService.js",
        "symptom": "TypeError: db.connect is not a function — server crashes on first admin request",
        "fix": "The database module exports { query, getClient, pool }. Use db.getClient() "
               "to get a pooled connection, or db.query() for direct queries. "
               "Never call db.connect() or pool.connect() directly.",
        "error_code": "EMI-SRV-001",
    },
    {
        "title": "HIGH — Admin API bypassed LockService (Phantom Locks)",
        "severity": "HIGH",
        "modules": "backend-admin-api",
        "files": "backend/src/modules/admin/adminDeviceService.js",
        "symptom": "Admin locks a device via dashboard — DB updates but phone does not actually lock",
        "fix": "All lock/unlock actions MUST go through LockService.requestLock() and "
               "LockService.requestUnlock(). Never call FirebaseService, AmapiService, or "
               "CommandSigningService directly from the admin module. LockService handles "
               "3-channel delivery, verification, and audit logging.",
        "error_code": "EMI-LOCK-303",
    },
    {
        "title": "HIGH — Triple redundant middleware in admin routes",
        "severity": "MEDIUM",
        "modules": "backend-admin-api",
        "files": "backend/src/modules/admin/adminRoutes.js",
        "symptom": "Conflicting middleware ordering causes some routes to return 403 unexpectedly",
        "fix": "Apply requireAuth, requireRole('admin'), and require2FA ONCE at router level. "
               "Do not repeat per-route. The router-level stack covers all routes underneath.",
        "error_code": "EMI-AUTH-105",
    },
    {
        "title": "HIGH — IMEI auth uses encrypted column for plaintext comparison",
        "severity": "HIGH",
        "modules": "backend-server-entry",
        "files": "backend/src/middleware/deviceAuth.js",
        "symptom": "Kotlin app sends plaintext IMEI in x-device-imei header — never matches "
                   "encrypted value in DB — all device auth fails",
        "fix": "Query by the plaintext imei column, not hardware_imei_encrypted. "
               "Encrypted columns are for storage protection, not for query matching.",
        "error_code": "EMI-DEV-201",
    },
    {
        "title": "HIGH — Redis operator precedence bug",
        "severity": "HIGH",
        "modules": "backend-server-entry",
        "files": "backend/src/config/redis.js",
        "symptom": "!redis.status === 'ready' always evaluates to false — Redis failures "
                   "silently ignored at startup",
        "fix": "Change to redis.status !== 'ready'. The ! operator has higher precedence "
               "than === so !redis.status converts to boolean first, then compares to string.",
        "error_code": "EMI-SRV-002",
    },
    {
        "title": "MEDIUM — backend-auth score 62 — weakest security module",
        "severity": "MEDIUM",
        "modules": "backend-auth",
        "files": "backend/src/modules/auth/jwt.js\nbackend/src/modules/auth/totp.js",
        "symptom": "Auth module has the lowest score of all modules. Risk of auth bypass "
                   "in production before it is fixed to 95+.",
        "fix": "Verify RS256 is used (not HS256). Verify refresh token rotation deletes old "
               "token in Redis. Verify TOTP backup codes are bcrypt-hashed. Verify rate "
               "limiter uses Redis store (not memory — memory resets on restart).",
        "error_code": "EMI-AUTH-001 / EMI-AUTH-104",
    },
    {
        "title": "MEDIUM — Decoupling: Dealer fraud flag must NOT block admin",
        "severity": "MEDIUM",
        "modules": "backend-decoupling",
        "files": "backend/src/modules/decoupling/decouplingService.js",
        "symptom": "If implemented incorrectly, a fraud flag could prevent executeDecouple() "
                   "from running, holding the device hostage indefinitely",
        "fix": "fraudFlag() sets state = FRAUD_FLAGGED but executeDecouple() must accept "
               "both PENDING_ADMIN_DECOUPLE and FRAUD_FLAGGED as valid states to proceed. "
               "Dealer flag = admin notification only, never a blocker.",
        "error_code": "EMI-DCP-500 / EMI-DCP-501",
    },
    {
        "title": "LOW — Antigravity supervisor MCP flag breaks Gemini CLI",
        "severity": "LOW",
        "modules": "build tooling (not production code)",
        "files": "gemini_fixer.py — run_gemini()",
        "symptom": "--allowed-mcp-server-names \"\" passes empty string to Gemini's policy "
                   "engine which rejects it: 'mcpName is required if specified'",
        "fix": "Remove --allowed-mcp-server-names flag entirely. Gemini handles MCP gracefully "
               "in yolo mode. The flag was intended to suppress noise but it causes a crash.",
        "error_code": "N/A (tooling only)",
    },
]

for p in problems:
    sev_color = C_RED if p["severity"] == "HIGH" else (C_AMBER if p["severity"] == "MEDIUM" else C_GREY)
    add_para(doc, f"[{p['severity']}]  {p['title']}", bold=True, color=sev_color, size=11)
    make_table(doc,
        ["Field", "Detail"],
        [
            [("Affected Module(s)", C_GREY, True), p["modules"]],
            [("File(s)",            C_GREY, True), p["files"]],
            [("Symptom",            C_GREY, True), p["symptom"]],
            [("Fix",                C_GREY, True), p["fix"]],
            [("Error Code",         C_GREY, True), (p["error_code"], C_BLUE, True)],
        ],
        col_widths=[1.2, 5.3], header_bg=C_GREY, stripe=False
    )
    doc.add_paragraph()

doc.add_page_break()

# ── SECTION 6: ADMIN PANEL DEBUG LOG INTEGRATION ─────────────────────────────
add_heading(doc, "6. Admin Panel Debug Log Integration", 1, C_BLUE)
add_para(doc,
    "The admin panel has a dedicated debug log tab. Every server error must emit a structured "
    "log event that the admin panel can display with the error code, module, severity, and "
    "a direct link to this blueprint section.", size=10)
doc.add_paragraph()

add_heading(doc, "6.1  Server-Side Error Emitter", 2)
add_para(doc, "Add this helper to backend/src/utils/errorEmitter.js:", size=9, italic=True, color=C_GREY)
code_lines = [
    "const emitDebugLog = (code, message, context = {}) => {",
    "  const entry = {",
    "    code,           // e.g. 'EMI-AUTH-101'",
    "    message,        // human-readable description",
    "    module: code.split('-')[1],",
    "    severity: getSeverity(code),",
    "    timestamp: new Date().toISOString(),",
    "    context,        // { userId, deviceId, requestId, ... }",
    "  };",
    "  logger.error(entry);",
    "  // Broadcast to admin panel via SSE or WebSocket",
    "  adminSSE.broadcast('debug_log', entry);",
    "};",
    "",
    "// Usage in any module:",
    "throw emitDebugLog('EMI-LOCK-303', 'All delivery channels failed', { deviceId });",
]
for line in code_lines:
    add_code(doc, "  " + line)

doc.add_paragraph()
add_heading(doc, "6.2  Admin Panel Debug Log Display", 2)
add_para(doc, "The debug log in the admin panel should show:", size=10)
make_table(doc,
    ["Column", "Source", "Purpose"],
    [
        ["Timestamp",  "entry.timestamp",       "When the error occurred"],
        ["Error Code", "entry.code",             "Click → opens this blueprint at section 4.3"],
        ["Module",     "entry.module",           "Which backend module threw the error"],
        ["Severity",   "entry.severity",         "HIGH (red) · MEDIUM (amber) · LOW (grey)"],
        ["Message",    "entry.message",          "Human-readable description"],
        ["Device ID",  "entry.context.deviceId", "Which device was affected (if applicable)"],
        ["Request ID", "entry.context.requestId","Correlates with server access log for full trace"],
    ],
    col_widths=[0.9, 1.6, 4.0], header_bg=C_BLACK
)

doc.add_paragraph()
add_heading(doc, "6.3  Severity Colour Mapping", 2)
make_table(doc,
    ["Code Range", "Severity", "Admin Panel Colour", "Action Required"],
    [
        ["001–099", "CRITICAL", ("Red background", C_RED, True),   "Immediate — server may be down"],
        ["100–199", "HIGH",     ("Red text",        C_RED, False),  "Same session — security risk"],
        ["200–299", "HIGH",     ("Red text",        C_RED, False),  "Same session — data integrity"],
        ["300–399", "MEDIUM",   ("Amber text",      C_AMBER, False),"Within 1 hour — service degraded"],
        ["400–499", "LOW",      ("Grey text",       C_GREY, False), "Log only — client input error"],
        ["500–599", "HIGH",     ("Red text",        C_RED, False),  "Immediate — state machine broken"],
        ["600–699", "CRITICAL", ("Red background",  C_RED, True),   "Immediate — security breach"],
        ["700–799", "LOW",      ("Amber text",      C_AMBER, False),"Monitor — potential abuse pattern"],
    ],
    col_widths=[0.9, 0.9, 1.5, 3.2], header_bg=C_BLACK
)

doc.add_page_break()

# ── SECTION 7: ENVIRONMENT VARIABLES ─────────────────────────────────────────
add_heading(doc, "7. Environment Variables Reference", 1, C_BLUE)
add_para(doc, "All URLs and secrets must come from environment variables. No hardcoded values anywhere.", size=10)
doc.add_paragraph()

make_table(doc,
    ["Variable", "Required", "Default (dev)", "Description"],
    [
        ["POSTGRES_URL",          ("YES", C_RED, True),   "postgresql://postgres:pass@localhost:5432/emilocker_dev", "PostgreSQL connection string"],
        ["REDIS_URL",             ("YES", C_RED, True),   "redis://localhost:6379",      "Redis connection string"],
        ["JWT_SECRET",            ("YES", C_RED, True),   "CHANGE_IN_PRODUCTION (≥32 chars)", "RS256 JWT signing secret"],
        ["HMAC_SECRET",           ("YES", C_RED, True),   "CHANGE_IN_PRODUCTION (≥32 chars)", "HMAC-SHA256 key signing secret"],
        ["FIREBASE_PROJECT_ID",   ("YES", C_RED, True),   "REPLACE_WITH_FIREBASE_ID",    "Firebase project for FCM"],
        ["AMAPI_PROJECT",         ("YES", C_RED, True),   "REPLACE_WITH_AMAPI_PROJECT",  "Google Android Management API project"],
        ["TWILIO_ACCOUNT_SID",    ("YES", C_RED, True),   "REPLACE_WITH_TWILIO_SID",     "Twilio account for SMS fallback"],
        ["TWILIO_AUTH_TOKEN",     ("YES", C_RED, True),   "REPLACE_WITH_TWILIO_TOKEN",   "Twilio auth token"],
        ["TWILIO_FROM_NUMBER",    ("YES", C_RED, True),   "+1XXXXXXXXXX",               "Twilio sender phone number"],
        ["API_BASE_URL",          ("YES", C_RED, True),   "http://localhost:3000",       "Backend API URL (used by all clients)"],
        ["VITE_API_BASE_URL",     ("YES", C_RED, True),   "http://localhost:3000",       "Admin panel API URL (Vite prefix required)"],
        ["CORS_ORIGIN",           ("YES", C_RED, True),   "http://localhost:5173",       "Allowed CORS origin for admin panel"],
        ["PORT",                  ("YES", C_RED, True),   "3000",                       "Express server port"],
        ["NODE_ENV",              ("YES", C_RED, True),   "development",                "Enables/disables SSL strict mode"],
        ["DEVICE_SIGNING_SECRET", ("YES", C_RED, True),   "CHANGE_IN_PRODUCTION",       "Signs device commands (separate from JWT)"],
        ["HARDWARE_BINDING_KEY",  ("YES", C_RED, True),   "CHANGE_IN_PRODUCTION",       "Hardware fingerprint binding key"],
    ],
    col_widths=[1.8, 0.7, 2.1, 2.0], header_bg=C_BLACK
)

doc.add_paragraph()
add_para(doc, "⚠  When deploying to production: replace ALL 'CHANGE_IN_PRODUCTION' and 'REPLACE_WITH_*' values before starting the server. "
    "The envValidator.js will refuse to start if any required variable is missing.", bold=True, color=C_RED, size=9)

doc.add_page_break()

# ── SECTION 8: SERVICE DEPENDENCY MAP ────────────────────────────────────────
add_heading(doc, "8. Service Dependency Map", 1, C_BLUE)
add_para(doc, "Who calls whom. Never bypass the layer below you — always call the service directly above the data layer.", size=10)
doc.add_paragraph()

dep_lines = [
    "  Controller (HTTP layer — thin, no business logic)",
    "       │  calls",
    "       ▼",
    "  Service (business logic — all rules live here)",
    "       │  calls",
    "       ├──► LockService          ← for any lock/unlock action",
    "       ├──► NotificationService  ← for any FCM/SMS send",
    "       ├──► AmapiService         ← for AMAPI operations",
    "       ├──► AuthService          ← for JWT/2FA checks",
    "       └──► Repository / Model   ← for DB read/write",
    "                  │  calls",
    "                  ▼",
    "          db.query() / db.getClient()   ← from database.js",
    "                  │",
    "                  ▼",
    "           PostgreSQL pool",
    "",
    "  ❌ NEVER DO THIS (causes phantom locks, bypassed auth):",
    "       AdminController → FirebaseService.send()    (must go through LockService)",
    "       AdminController → AmapiService.applyPolicy() (must go through LockService)",
    "       EmiController   → decouplingService directly (must go through EmiService)",
]
for line in dep_lines:
    add_code(doc, line)

doc.add_page_break()

# ── SECTION 9: AI ORCHESTRATION ──────────────────────────────────────────────
add_heading(doc, "9. AI Build Orchestration", 1, C_BLUE)
add_para(doc, "This codebase was built autonomously by a 4-AI orchestration system. "
    "Understanding the build pipeline helps when debugging generated code.", size=10)
doc.add_paragraph()

make_table(doc,
    ["AI Agent", "Role", "Model", "Trigger"],
    [
        [("MIMO V2.5 Pro", C_BLUE, True),    "Executor / Reviewer",  "xiaomi-token-plan-singapore/mimo-v2.5-pro (1M ctx)", "Reviews every module, assigns score 0–100"],
        [("MiniMax M2.7",  C_VIOLET, True),  "Worker / Implementer", "minimax-coding-plan/MiniMax-M2.7 (256k ctx)",         "Implements modules, applies fix rounds"],
        [("Gemini CLI",    C_GREEN, True),   "First Fixer",          "Gemini CLI v0.40.1 (Antigravity)",                    "Fixes under-95 modules — up to 5 rounds"],
        [("Codex ZGBT 5.5",C_CYAN, True),   "Second Fixer",         "gpt-5.4 via Codex CLI",                               "Activated if Gemini fails all 5 rounds"],
        [("Claude Sonnet", C_AMBER, True),  "Heavy Lifter",         "claude-sonnet-4-6",                                   "Activated if Codex also fails — writes new fix plan"],
    ],
    col_widths=[1.2, 1.3, 2.3, 2.2], header_bg=C_BLACK
)

doc.add_paragraph()
add_para(doc, "Quality Gate: A module is only marked complete when MIMO gives status=approved AND score ≥ 95 AND zero HIGH severity issues.", bold=True, color=C_GREEN, size=10)

# ── FOOTER ────────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "Document End", 2, C_GREY)
add_para(doc, f"EMI Locker Platform Blueprint v1.0 — Generated {datetime.date.today()}", color=C_GREY, size=9)
add_para(doc, "Keep this document updated as new modules are added or bugs are discovered.", italic=True, color=C_GREY, size=9)
add_para(doc, "Error codes in this document must match exactly the codes thrown in the production codebase.", italic=True, color=C_RED, size=9)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out_path = r"d:\EMI APP\EMI_Locker_Blueprint.docx"
doc.save(out_path)
print(f"Blueprint saved: {out_path}")
